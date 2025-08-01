import os
import fitz  # PyMuPDF
from docx import Document
from typing import List, Dict, Optional, Tuple
from pathlib import Path
import hashlib
from loguru import logger
from config.settings import settings

class DocumentParser:
    """Service for parsing various document formats (PDF, DOCX, TXT)."""
    
    def __init__(self):
        """Initialize document parser."""
        self.supported_formats = settings.FILE_SETTINGS["supported_formats"]
        self.max_file_size = settings.FILE_SETTINGS["max_file_size_mb"] * 1024 * 1024
    
    def process_multiple_cvs(self, cv_files_list: List) -> List[Dict]:
        """
        Process multiple CV files and extract text content.
        
        Args:
            cv_files_list: List of file objects (UploadFile, file-like objects, or file paths)
            
        Returns:
            List of dictionaries containing parsed CV data
        """
        processed_cvs = []
        
        for file_obj in cv_files_list:
            try:
                cv_data = self._process_single_file(file_obj, "cv")
                if cv_data:
                    processed_cvs.append(cv_data)
            except Exception as e:
                logger.error(f"Error processing CV file: {str(e)}")
                continue
        
        logger.info(f"Successfully processed {len(processed_cvs)} CV files")
        return processed_cvs
    
    def process_multiple_jobs(self, job_files_list: List) -> List[Dict]:
        """
        Process multiple job description files and extract text content.
        
        Args:
            job_files_list: List of file objects (UploadFile, file-like objects, or file paths)
            
        Returns:
            List of dictionaries containing parsed job data
        """
        processed_jobs = []
        
        for file_obj in job_files_list:
            try:
                job_data = self._process_single_file(file_obj, "job")
                if job_data:
                    processed_jobs.append(job_data)
            except Exception as e:
                logger.error(f"Error processing job file: {str(e)}")
                continue
        
        logger.info(f"Successfully processed {len(processed_jobs)} job files")
        return processed_jobs
    
    def _process_single_document(self, file_path: str, doc_type: str) -> Optional[Dict]:
        """
        Process a single document and extract its content.
        
        Args:
            file_path: Path to the document file
            doc_type: Type of document ("cv" or "job")
            
        Returns:
            Dictionary containing parsed document data
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            logger.error(f"File too large: {file_path} ({file_size} bytes)")
            return None
        
        # Check file format
        file_extension = Path(file_path).suffix.lower()
        if file_extension not in self.supported_formats:
            logger.error(f"Unsupported file format: {file_extension}")
            return None
        
        try:
            # Extract text based on file format
            if file_extension == ".pdf":
                text_content = self._extract_pdf_text(file_path)
            elif file_extension == ".docx":
                text_content = self._extract_docx_text(file_path)
            elif file_extension == ".txt":
                text_content = self._extract_txt_text(file_path)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            if not text_content or not text_content.strip():
                logger.warning(f"No text content extracted from: {file_path}")
                return None
            
            # Generate document ID
            doc_id = self._generate_document_id(file_path, text_content)
            
            # Create document data structure
            document_data = {
                "doc_id": doc_id,
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_size": file_size,
                "file_extension": file_extension,
                "doc_type": doc_type,
                "raw_text": text_content,
                "processed_at": None  # Will be set by the processor
            }
            
            logger.info(f"Successfully processed {doc_type} document: {file_path}")
            return document_data
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return None
    
    def _process_single_file(self, file_obj, doc_type: str) -> Optional[Dict]:
        """
        Process a single file object and extract its content.
        
        Args:
            file_obj: File object (UploadFile, file-like object, or file path)
            doc_type: Type of document ("cv" or "job")
            
        Returns:
            Dictionary containing parsed document data
        """
        try:
            # Handle different types of file objects
            if hasattr(file_obj, 'filename') and hasattr(file_obj, 'file'):
                # FastAPI UploadFile object
                file_name = file_obj.filename
                file_size = file_obj.size if hasattr(file_obj, 'size') else 0
                file_extension = Path(file_name).suffix.lower()
                file_content = file_obj.file.read()
                
                # Check file size
                if file_size > self.max_file_size:
                    logger.error(f"File too large: {file_name} ({file_size} bytes)")
                    return None
                
            elif hasattr(file_obj, 'name') and hasattr(file_obj, 'read'):
                # File-like object
                file_name = getattr(file_obj, 'name', 'unknown')
                file_content = file_obj.read()
                file_size = len(file_content)
                file_extension = Path(file_name).suffix.lower()
                
                # Check file size
                if file_size > self.max_file_size:
                    logger.error(f"File too large: {file_name} ({file_size} bytes)")
                    return None
                
            elif isinstance(file_obj, str):
                # File path string - use the existing method
                return self._process_single_document(file_obj, doc_type)
            else:
                logger.error(f"Unsupported file object type: {type(file_obj)}")
                return None
            
            # Check file format
            if file_extension not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            # Extract text based on file format
            if file_extension == ".pdf":
                text_content = self._extract_pdf_from_bytes(file_content)
            elif file_extension == ".docx":
                text_content = self._extract_docx_from_bytes(file_content)
            elif file_extension == ".txt":
                text_content = self._extract_txt_from_bytes(file_content)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            if not text_content or not text_content.strip():
                logger.warning(f"No text content extracted from: {file_name}")
                return None
            
            # Generate document ID
            doc_id = self._generate_document_id_from_content(file_name, text_content)
            
            # Create document data structure
            document_data = {
                "doc_id": doc_id,
                "file_path": None,  # No file path for uploaded files
                "file_name": file_name,
                "file_size": file_size,
                "file_extension": file_extension,
                "doc_type": doc_type,
                "raw_text": text_content,
                "processed_at": None  # Will be set by the processor
            }
            
            logger.info(f"Successfully processed {doc_type} document: {file_name}")
            return document_data
            
        except Exception as e:
            logger.error(f"Error processing file object: {str(e)}")
            return None
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyMuPDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            doc = fitz.open(file_path)
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content).strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {str(e)}")
            raise
    
    def _extract_txt_text(self, file_path: str) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        return content.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                return content.strip()
                
        except Exception as e:
            logger.error(f"Error extracting TXT text from {file_path}: {str(e)}")
            raise
    
    def _generate_document_id(self, file_path: str, content: str) -> str:
        """
        Generate a unique document ID based on file path and content.
        
        Args:
            file_path: Path to the document
            content: Document content
            
        Returns:
            Unique document ID
        """
        # Create a hash based on file path and content
        hash_input = f"{file_path}:{content[:1000]}"  # Use first 1000 chars for performance
        return hashlib.md5(hash_input.encode()).hexdigest()
    
    def validate_document(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate if a document can be processed.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            return False, f"File too large ({file_size} bytes > {self.max_file_size} bytes)"
        
        file_extension = Path(file_path).suffix.lower()
        if file_extension not in self.supported_formats:
            return False, f"Unsupported file format: {file_extension}"
        
        return True, "Valid document"
    
    def get_document_metadata(self, file_path: str) -> Dict:
        """
        Get metadata about a document without processing its content.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing document metadata
        """
        if not os.path.exists(file_path):
            return {}
        
        stat = os.stat(file_path)
        return {
            "file_name": Path(file_path).name,
            "file_size": stat.st_size,
            "file_extension": Path(file_path).suffix.lower(),
            "created_time": stat.st_ctime,
            "modified_time": stat.st_mtime,
            "is_readable": os.access(file_path, os.R_OK)
        }

    def _extract_pdf_from_bytes(self, file_content: bytes) -> str:
        """
        Extract text from PDF bytes using PyMuPDF.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            import io
            doc = fitz.open(stream=io.BytesIO(file_content), filetype="pdf")
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text from bytes: {str(e)}")
            raise
    
    def _extract_docx_from_bytes(self, file_content: bytes) -> str:
        """
        Extract text from DOCX bytes using python-docx.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            import io
            doc = Document(io.BytesIO(file_content))
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content).strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text from bytes: {str(e)}")
            raise
    
    def _extract_txt_from_bytes(self, file_content: bytes) -> str:
        """
        Extract text from TXT bytes.
        
        Args:
            file_content: TXT file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    content = file_content.decode(encoding)
                    return content.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            content = file_content.decode('utf-8', errors='ignore')
            return content.strip()
                
        except Exception as e:
            logger.error(f"Error extracting TXT text from bytes: {str(e)}")
            raise
    
    def _generate_document_id_from_content(self, file_name: str, content: str) -> str:
        """
        Generate a unique document ID based on file name and content.
        
        Args:
            file_name: Name of the file
            content: Document content
            
        Returns:
            Unique document ID
        """
        # Create a hash based on file name and content
        hash_input = f"{file_name}:{content[:1000]}"  # Use first 1000 chars for performance
        return hashlib.md5(hash_input.encode()).hexdigest()

# Global document parser instance
document_parser = DocumentParser() 